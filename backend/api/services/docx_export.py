"""DOCX exporters for benchmark receipts and chat sessions.

Lawyers live in Word. Markdown export is for devs; this module emits the
minimum-viable .docx for the legal user. Footer auto-injects the README
disclaimer on every export.
"""
from __future__ import annotations
import io
import re
from datetime import datetime, timezone
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

DISCLAIMER_TITLE = "For Informational Purposes Only"
DISCLAIMER_BODY = (
    "The information provided on Junas is for general informational purposes only. "
    "While we strive to ensure the accuracy and reliability of the legal analysis "
    "workflows and templates provided, Junas makes no guarantees, representations, "
    "or warranties of any kind, express or implied, about the completeness, accuracy, "
    "reliability, suitability, or availability of the information. Users should "
    "independently verify any information before making decisions based on it. "
    "Junas does not provide professional legal advice or consultation services."
)  # mirrors README.md §For Informational Purposes Only + No Professional Advice

_SLUG_RE = re.compile(r"[^a-z0-9]+")
_FENCE_RE = re.compile(r"^```(\w*)\s*$")
_TABLE_ROW_RE = re.compile(r"^\s*\|.+\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?[\s\-:|]+\|?\s*$")
_BULLET_RE = re.compile(r"^(\s*)([-*+])\s+(.*)$")
_NUMLIST_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_INLINE_CODE_RE = re.compile(r"`([^`]+)`")
_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")


def slugify(value: str, max_len: int = 48) -> str:
    lowered = (value or "").strip().lower()
    slug = _SLUG_RE.sub("-", lowered).strip("-")
    if not slug:
        return "untitled"
    return slug[:max_len].rstrip("-") or "untitled"


def _set_courier(run, size: int = 9) -> None:
    run.font.name = "Courier New"
    run.font.size = Pt(size)


def _append_disclaimer(doc: DocxDocument) -> None:
    # footer-style block on the last page; python-docx footers are per-section
    # and don't render well for long content, so we append at body end with a
    # visual separator. lawyers expect to see it inline anyway.
    doc.add_paragraph().add_run().add_break()
    sep = doc.add_paragraph()
    sep_run = sep.add_run("—" * 40)
    sep_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    title = doc.add_paragraph()
    title_run = title.add_run(DISCLAIMER_TITLE)
    title_run.bold = True
    title_run.font.size = Pt(9)
    body = doc.add_paragraph()
    body_run = body.add_run(DISCLAIMER_BODY)
    body_run.font.size = Pt(8)
    body_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _add_inline(paragraph, text: str) -> None:
    # split on inline code first (highest precedence), then bold, then italic.
    # python-docx has no rich text parser; we tokenise by regex and emit runs.
    if not text:
        return
    pos = 0
    tokens: list[tuple[str, str]] = []  # (kind, content)
    for m in _INLINE_CODE_RE.finditer(text):
        if m.start() > pos:
            tokens.append(("text", text[pos:m.start()]))
        tokens.append(("code", m.group(1)))
        pos = m.end()
    if pos < len(text):
        tokens.append(("text", text[pos:]))

    def _emit_text(chunk: str) -> None:
        # bold/italic pass; supports either, not nested.
        cursor = 0
        for bm in _BOLD_RE.finditer(chunk):
            if bm.start() > cursor:
                _emit_italic(chunk[cursor:bm.start()])
            run = paragraph.add_run(bm.group(1))
            run.bold = True
            cursor = bm.end()
        if cursor < len(chunk):
            _emit_italic(chunk[cursor:])

    def _emit_italic(chunk: str) -> None:
        cursor = 0
        for im in _ITALIC_RE.finditer(chunk):
            if im.start() > cursor:
                paragraph.add_run(chunk[cursor:im.start()])
            run = paragraph.add_run(im.group(1))
            run.italic = True
            cursor = im.end()
        if cursor < len(chunk):
            paragraph.add_run(chunk[cursor:])

    for kind, content in tokens:
        if kind == "code":
            run = paragraph.add_run(content)
            _set_courier(run)
        else:
            _emit_text(content)


def _add_markdown(doc: DocxDocument, markdown: str) -> None:
    # block-level walker: headings, fenced code, pipe tables, bullet/numbered
    # lists, blank lines as paragraph breaks. anything else → plain paragraph.
    lines = (markdown or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        # fenced code
        fence_match = _FENCE_RE.match(line)
        if fence_match:
            buf: list[str] = []
            i += 1
            while i < len(lines) and not _FENCE_RE.match(lines[i]):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            para = doc.add_paragraph()
            run = para.add_run("\n".join(buf))
            _set_courier(run)
            continue
        # heading
        h = _HEADING_RE.match(line)
        if h:
            level = min(len(h.group(1)), 4)
            doc.add_heading(h.group(2).strip(), level=level)
            i += 1
            continue
        # table
        if _TABLE_ROW_RE.match(line) and i + 1 < len(lines) and _TABLE_SEP_RE.match(lines[i + 1]):
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip separator
            body_rows: list[list[str]] = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i]):
                body_rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1 + len(body_rows), cols=len(header_cells))
            table.style = "Light Grid"
            for col, cell_text in enumerate(header_cells):
                cell = table.rows[0].cells[col]
                cell.text = ""
                run = cell.paragraphs[0].add_run(cell_text)
                run.bold = True
            for r, row in enumerate(body_rows, start=1):
                for c, cell_text in enumerate(row[:len(header_cells)]):
                    table.rows[r].cells[c].text = cell_text
            continue
        # bullets
        bm = _BULLET_RE.match(line)
        if bm:
            indent_level = len(bm.group(1)) // 2
            para = doc.add_paragraph(style="List Bullet" if indent_level == 0 else "List Bullet 2")
            _add_inline(para, bm.group(3))
            i += 1
            continue
        # numbered
        nm = _NUMLIST_RE.match(line)
        if nm:
            indent_level = len(nm.group(1)) // 2
            para = doc.add_paragraph(style="List Number" if indent_level == 0 else "List Number 2")
            _add_inline(para, nm.group(3))
            i += 1
            continue
        # blank
        if not line.strip():
            i += 1
            continue
        # default paragraph
        para = doc.add_paragraph()
        _add_inline(para, line)
        i += 1


def _format_dt(value: str | None) -> str:
    if not value:
        return "-"
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).strftime("%Y-%m-%d %H:%M UTC")
    except ValueError:
        return value


def build_receipt_docx(receipt: dict[str, Any]) -> bytes:
    doc = Document()
    # header
    title = doc.add_heading("SG-LegalBench Receipt", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    workflow = str(receipt.get("workflow") or "-")
    provenance = receipt.get("provenance") or {}
    model = str(provenance.get("model_display_name") or provenance.get("model") or "-")
    finished = _format_dt(receipt.get("finished_at"))
    meta = doc.add_paragraph()
    meta.add_run("Task: ").bold = True
    meta.add_run(workflow)
    meta.add_run("    Model: ").bold = True
    meta.add_run(model)
    meta.add_run("    Finished: ").bold = True
    meta.add_run(finished)
    dataset = str(receipt.get("dataset") or "-")
    p = doc.add_paragraph()
    p.add_run("Dataset: ").bold = True
    p.add_run(dataset)
    p = doc.add_paragraph()
    p.add_run("Cases: ").bold = True
    p.add_run(str(receipt.get("total_cases") or 0))
    p.add_run("    Strict: ").bold = True
    p.add_run("yes" if receipt.get("strict") else "no")
    p.add_run("    Tier: ").bold = True
    p.add_run(str(receipt.get("data_tier") or "regulator"))

    # per-evaluator means
    doc.add_heading("Per-evaluator means", level=1)
    means = receipt.get("per_evaluator_mean") or {}
    if means:
        table = doc.add_table(rows=1 + len(means), cols=2)
        table.style = "Light Grid"
        hdr = table.rows[0].cells
        hdr[0].paragraphs[0].add_run("Evaluator").bold = True
        hdr[1].paragraphs[0].add_run("Mean score").bold = True
        for r, (name, score) in enumerate(sorted(means.items()), start=1):
            table.rows[r].cells[0].text = str(name)
            try:
                table.rows[r].cells[1].text = f"{float(score):.3f}"
            except (TypeError, ValueError):
                table.rows[r].cells[1].text = str(score)
    else:
        doc.add_paragraph("(no evaluator means recorded)")

    # per-case table
    doc.add_heading("Per-case results", level=1)
    results = [r for r in (receipt.get("results") or []) if isinstance(r, dict)]
    by_case: dict[str, dict[str, Any]] = {}
    evaluator_names: list[str] = []
    seen_evals: set[str] = set()
    for r in results:
        case_name = str(r.get("case_name") or "")
        evaluator = str(r.get("evaluator") or "")
        if not case_name or not evaluator:
            continue
        if evaluator not in seen_evals:
            seen_evals.add(evaluator)
            evaluator_names.append(evaluator)
        by_case.setdefault(case_name, {})[evaluator] = r.get("score")
    if by_case:
        table = doc.add_table(rows=1 + len(by_case), cols=1 + len(evaluator_names))
        table.style = "Light Grid"
        hdr = table.rows[0].cells
        hdr[0].paragraphs[0].add_run("Case").bold = True
        for ci, ev in enumerate(evaluator_names, start=1):
            hdr[ci].paragraphs[0].add_run(ev).bold = True
        for ri, (case_name, scores) in enumerate(sorted(by_case.items()), start=1):
            table.rows[ri].cells[0].text = case_name
            for ci, ev in enumerate(evaluator_names, start=1):
                score = scores.get(ev)
                if isinstance(score, (int, float)):
                    table.rows[ri].cells[ci].text = f"{float(score):.3f}"
                else:
                    table.rows[ri].cells[ci].text = "-"
    else:
        doc.add_paragraph("(no per-case results recorded)")

    _append_disclaimer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_session_docx(session: dict[str, Any]) -> bytes:
    doc = Document()
    title = str(session.get("title") or "Chat session")
    doc.add_heading(title, level=0)
    created = _format_dt(session.get("created_at"))
    meta = doc.add_paragraph()
    meta.add_run("Exported: ").bold = True
    meta.add_run(datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))
    if created and created != "-":
        meta.add_run("    Created: ").bold = True
        meta.add_run(created)

    messages: Iterable[dict[str, Any]] = session.get("messages") or []
    for msg in messages:
        role = str(msg.get("role") or "message").lower()
        ts = msg.get("timestamp")
        ts_label = ""
        if isinstance(ts, (int, float)):
            try:
                seconds = float(ts) / 1000 if ts > 1e12 else float(ts)
                ts_label = datetime.fromtimestamp(seconds, tz=timezone.utc).strftime("%H:%M:%S")
            except (OSError, ValueError):
                ts_label = ""
        elif isinstance(ts, str) and ts:
            ts_label = _format_dt(ts)
        header_para = doc.add_paragraph()
        role_label = "You" if role == "user" else ("Junas" if role == "assistant" else role.title())
        role_run = header_para.add_run(role_label)
        role_run.bold = True
        if ts_label:
            ts_run = header_para.add_run(f"  {ts_label}")
            ts_run.font.size = Pt(9)
            ts_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        _add_markdown(doc, str(msg.get("content") or ""))

    _append_disclaimer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def receipt_filename(run_id: str) -> str:
    safe_run = slugify(run_id, max_len=80) or "run"
    return f"junas-receipt-{safe_run}.docx"


def session_filename(session_id: str, title: str | None) -> str:
    slug = slugify(title or "", max_len=48)
    return f"junas-session-{slugify(session_id, max_len=32)}-{slug}.docx"

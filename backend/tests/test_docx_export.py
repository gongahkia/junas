"""Tests for DOCX export service + /exports/* endpoints.

Covers the receipt path against a real fixture from runs/baselines/. The
session export path is exercised via inline payload (COPILOT-1 not landed).
"""
from __future__ import annotations
import io
import json
import time
import zipfile
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from api.main import create_app, settings
from api.services.docx_export import (
    DISCLAIMER_TITLE,
    build_receipt_docx,
    build_session_docx,
    receipt_filename,
    session_filename,
    slugify,
)


AUTH_HEADERS = {"X-API-Key": "test-key"}
REPO_ROOT = Path(__file__).resolve().parents[2]
BASELINE_RECEIPT = (
    REPO_ROOT / "runs" / "baselines" / "ollama" / "sglb_01" / "20260604T073720Z-qwen2-5vl-7b.json"
)


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _is_valid_docx(blob: bytes) -> bool:
    # docx is a zip with [Content_Types].xml at the root
    if not blob.startswith(b"PK"):
        return False
    with zipfile.ZipFile(io.BytesIO(blob)) as zf:
        return "[Content_Types].xml" in zf.namelist()


@pytest.fixture
def fixture_receipt() -> dict:
    assert BASELINE_RECEIPT.exists(), f"missing baseline fixture: {BASELINE_RECEIPT}"
    return json.loads(BASELINE_RECEIPT.read_text(encoding="utf-8"))


@pytest.fixture
def client(tmp_path, monkeypatch) -> TestClient:
    runs_dir = tmp_path / "runs"
    (runs_dir / "baselines" / "ollama" / "sglb_01").mkdir(parents=True)
    target = runs_dir / "baselines" / "ollama" / "sglb_01" / "20260604T073720Z-qwen2-5vl-7b.json"
    target.write_bytes(BASELINE_RECEIPT.read_bytes())
    monkeypatch.setenv("JUNAS_BENCHMARK_RUNS_DIR", str(runs_dir))
    monkeypatch.setattr(settings, "api_keys", ["test-key"])
    monkeypatch.setattr(settings, "require_auth", False)
    app = create_app()
    c = TestClient(app)
    c.headers.update(AUTH_HEADERS)
    return c


def test_slugify_strips_punctuation() -> None:
    assert slugify("PDPA outcome: First Pass!") == "pdpa-outcome-first-pass"
    assert slugify("") == "untitled"
    assert slugify("   ") == "untitled"


def test_receipt_filename_pattern() -> None:
    assert receipt_filename("abc__sglb_01__run") == "junas-receipt-abc-sglb-01-run.docx"


def test_session_filename_pattern() -> None:
    assert session_filename("conv_123", "My Tenancy Q&A") == "junas-session-conv-123-my-tenancy-q-a.docx"


def test_build_receipt_docx_writes_valid_file(tmp_path: Path, fixture_receipt: dict) -> None:
    blob = build_receipt_docx(fixture_receipt)
    out = tmp_path / "receipt.docx"
    out.write_bytes(blob)
    assert out.stat().st_size > 2000  # non-trivial output
    assert _is_valid_docx(blob)


def test_build_receipt_docx_contains_header_and_means(fixture_receipt: dict) -> None:
    blob = build_receipt_docx(fixture_receipt)
    text = _docx_text(blob)
    assert "SG-LegalBench Receipt" in text
    assert "sglb_01" in text  # workflow
    assert "qwen" in text.lower()  # model
    assert "sglb_01_obligations_f1" in text  # evaluator
    assert "penalty_band_mae" in text


def test_build_receipt_docx_includes_per_case_table(fixture_receipt: dict) -> None:
    blob = build_receipt_docx(fixture_receipt)
    text = _docx_text(blob)
    # at least one case_name from the fixture
    case_names = {r.get("case_name") for r in fixture_receipt["results"]}
    assert any(name in text for name in case_names if name)


def test_disclaimer_present_on_receipt(fixture_receipt: dict) -> None:
    blob = build_receipt_docx(fixture_receipt)
    text = _docx_text(blob)
    assert DISCLAIMER_TITLE in text
    assert "informational purposes only" in text.lower()


def test_session_docx_round_trips_code_and_table() -> None:
    session = {
        "title": "Tenancy Q&A",
        "messages": [
            {"role": "user", "content": "What does s 23 PDPA say?"},
            {
                "role": "assistant",
                "content": (
                    "# Heading\n\n"
                    "Some **bold** and *italic* and `inline_code`.\n\n"
                    "- bullet one\n- bullet two\n\n"
                    "| col1 | col2 |\n| --- | --- |\n| a | b |\n| c | d |\n\n"
                    "```python\nprint('hi')\n```\n"
                ),
                "timestamp": 1717000000000,
            },
        ],
    }
    blob = build_session_docx(session)
    assert _is_valid_docx(blob)
    text = _docx_text(blob)
    assert "Tenancy Q&A" in text
    assert "bullet one" in text
    assert "col1" in text and "col2" in text
    assert "print('hi')" in text
    assert "Heading" in text
    assert DISCLAIMER_TITLE in text


def test_session_docx_performance_under_3s() -> None:
    # 200 messages alternating roles, mixed markdown content
    messages = []
    for i in range(200):
        role = "user" if i % 2 == 0 else "assistant"
        content = (
            f"Message {i}: discussion of clause **{i}**.\n\n"
            "- bullet a\n- bullet b\n\n"
            f"```\ncode block {i}\n```\n"
        )
        messages.append({"role": role, "content": content, "timestamp": 1717000000 + i})
    session = {"title": "Perf test", "messages": messages}
    start = time.perf_counter()
    blob = build_session_docx(session)
    elapsed = time.perf_counter() - start
    assert _is_valid_docx(blob)
    assert elapsed < 3.0, f"export took {elapsed:.2f}s (limit 3.0s)"


def test_export_receipt_endpoint_returns_docx(client: TestClient) -> None:
    resp = client.get("/api/v1/exports/receipt/ollama__sglb_01__20260604T073720Z-qwen2-5vl-7b.docx")
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in resp.headers["content-disposition"]
    assert "junas-receipt-" in resp.headers["content-disposition"]
    assert _is_valid_docx(resp.content)


def test_export_receipt_unknown_run_404(client: TestClient) -> None:
    resp = client.get("/api/v1/exports/receipt/does-not-exist.docx")
    assert resp.status_code == 404


def test_export_session_without_body_returns_501(client: TestClient) -> None:
    resp = client.post("/api/v1/exports/session/conv_1.docx")
    assert resp.status_code == 501
    assert "not yet wired" in resp.json()["detail"]


def test_export_session_with_payload_returns_docx(client: TestClient) -> None:
    body = {
        "session_id": "conv_42",
        "title": "Quick chat",
        "messages": [
            {"role": "user", "content": "hello", "timestamp": 1717000000000},
            {"role": "assistant", "content": "hi there\n\n```\ncode\n```", "timestamp": 1717000005000},
        ],
    }
    resp = client.post("/api/v1/exports/session/conv_42.docx", json=body)
    assert resp.status_code == 200, resp.text
    assert _is_valid_docx(resp.content)
    assert "junas-session-conv-42-quick-chat.docx" in resp.headers["content-disposition"]


def test_export_session_id_mismatch_400(client: TestClient) -> None:
    body = {"session_id": "conv_999", "messages": [{"role": "user", "content": "x"}]}
    resp = client.post("/api/v1/exports/session/conv_42.docx", json=body)
    assert resp.status_code == 400
